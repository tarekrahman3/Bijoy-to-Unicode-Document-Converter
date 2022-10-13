from docx import Document
from bijoytounicode import bijoy2unicode


def recursiveCall(table):
	tableRows = table.rows
	for row in tableRows:
		rowCells =  row.cells
		for cell in rowCells:
			oldCellText = cell.text
			print(oldCellText)
			convertedCellText = bijoy2unicode(cell.text)
			print(convertedCellText)
			cell.text = convertedCellText
			for table in cell.tables:
				recursiveCall(table)


root = Document('inputDocumentFile.docx')

[recursiveCall(table) for table in root.tables]

root.save('OutputDocumentFile.docx')
